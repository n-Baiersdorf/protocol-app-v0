"""
File Service - Datei-Upload und -Management
"""

import os
import uuid
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional
from werkzeug.utils import secure_filename
from werkzeug.datastructures import FileStorage
import logging

logger = logging.getLogger(__name__)

class FileService:
    """Service für Datei-Upload und -Management"""
    
    ALLOWED_EXTENSIONS = {
        'images': {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff'},
        'documents': {'pdf', 'doc', 'docx', 'txt', 'rtf'},
        'data': {'csv', 'xlsx', 'xls', 'json'},
        'other': {'zip', 'rar'}
    }
    
    MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB
    
    def __init__(self, upload_folder: str):
        self.upload_folder = Path(upload_folder)
        self.upload_folder.mkdir(exist_ok=True)
        
        # Unterordner für verschiedene Dateitypen erstellen
        for category in self.ALLOWED_EXTENSIONS.keys():
            (self.upload_folder / category).mkdir(exist_ok=True)
    
    def save_uploaded_file(self, file: FileStorage) -> Dict:
        """
        Speichert eine hochgeladene Datei
        
        Args:
            file: Werkzeug FileStorage Objekt
            
        Returns:
            Dict mit Datei-Informationen
        """
        try:
            # Dateivalidierung
            if not file or file.filename == '':
                raise ValueError("Keine gültige Datei")
            
            if not self._is_allowed_file(file.filename):
                raise ValueError(f"Dateityp nicht erlaubt: {file.filename}")
            
            # Sicherer Dateiname generieren
            original_filename = secure_filename(file.filename)
            file_extension = self._get_file_extension(original_filename)
            unique_filename = f"{uuid.uuid4().hex}_{original_filename}"
            
            # Dateikategorie bestimmen
            file_category = self._get_file_category(file_extension)
            
            # Speicherpfad bestimmen
            save_path = self.upload_folder / file_category / unique_filename
            
            # Datei speichern
            file.save(save_path)
            
            # Datei-Informationen sammeln
            file_info = {
                'id': str(uuid.uuid4()),
                'original_name': original_filename,
                'filename': unique_filename,
                'path': str(save_path),
                'size': save_path.stat().st_size,
                'type': file_category,
                'extension': file_extension,
                'mime_type': mimetypes.guess_type(str(save_path))[0],
                'created_at': save_path.stat().st_ctime
            }
            
            # Zusätzliche Verarbeitung je nach Dateityp
            if file_category == 'images':
                file_info.update(self._process_image_file(save_path))
            elif file_category == 'documents':
                file_info.update(self._process_document_file(save_path))
            elif file_category == 'data':
                file_info.update(self._process_data_file(save_path))
            
            logger.info(f"Datei erfolgreich gespeichert: {unique_filename}")
            return file_info
            
        except Exception as e:
            logger.error(f"Fehler beim Speichern der Datei: {str(e)}")
            raise
    
    def _is_allowed_file(self, filename: str) -> bool:
        """Prüft, ob der Dateityp erlaubt ist"""
        if not filename:
            return False
        
        extension = self._get_file_extension(filename)
        all_extensions = set()
        for ext_set in self.ALLOWED_EXTENSIONS.values():
            all_extensions.update(ext_set)
        
        return extension in all_extensions
    
    def _get_file_extension(self, filename: str) -> str:
        """Extrahiert die Dateierweiterung"""
        return filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    def _get_file_category(self, extension: str) -> str:
        """Bestimmt die Dateikategorie basierend auf der Erweiterung"""
        for category, extensions in self.ALLOWED_EXTENSIONS.items():
            if extension in extensions:
                return category
        return 'other'
    
    def _process_image_file(self, file_path: Path) -> Dict:
        """Verarbeitet Bilddateien"""
        try:
            from PIL import Image
            
            with Image.open(file_path) as img:
                return {
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode
                }
        except Exception as e:
            logger.warning(f"Fehler bei der Bildverarbeitung: {str(e)}")
            return {'processing_error': str(e)}
    
    def _process_document_file(self, file_path: Path) -> Dict:
        """Verarbeitet Dokumentdateien"""
        try:
            info = {'pages': None, 'text_length': None}
            
            if file_path.suffix.lower() == '.pdf':
                info.update(self._process_pdf(file_path))
            elif file_path.suffix.lower() in ['.doc', '.docx']:
                info.update(self._process_word_doc(file_path))
            elif file_path.suffix.lower() == '.txt':
                info.update(self._process_text_file(file_path))
            
            return info
            
        except Exception as e:
            logger.warning(f"Fehler bei der Dokumentverarbeitung: {str(e)}")
            return {'processing_error': str(e)}
    
    def _process_pdf(self, file_path: Path) -> Dict:
        """Verarbeitet PDF-Dateien"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                return {
                    'pages': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else ''
                }
        except Exception as e:
            logger.warning(f"PDF-Verarbeitung fehlgeschlagen: {str(e)}")
            return {'pdf_error': str(e)}
    
    def _process_word_doc(self, file_path: Path) -> Dict:
        """Verarbeitet Word-Dokumente"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            return {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'images': len(doc.inline_shapes)
            }
        except Exception as e:
            logger.warning(f"Word-Verarbeitung fehlgeschlagen: {str(e)}")
            return {'word_error': str(e)}
    
    def _process_text_file(self, file_path: Path) -> Dict:
        """Verarbeitet Textdateien"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
            return {
                'text_length': len(content),
                'lines': len(content.splitlines()),
                'words': len(content.split())
            }
        except Exception as e:
            logger.warning(f"Text-Verarbeitung fehlgeschlagen: {str(e)}")
            return {'text_error': str(e)}
    
    def _process_data_file(self, file_path: Path) -> Dict:
        """Verarbeitet Datendateien (CSV, Excel)"""
        try:
            import pandas as pd
            
            info = {}
            
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
                info = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist()
                }
            elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
                info = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist(),
                    'sheets': 1  # Vereinfacht für MVP
                }
            
            return info
            
        except Exception as e:
            logger.warning(f"Datenverarbeitung fehlgeschlagen: {str(e)}")
            return {'data_error': str(e)}
    
    def get_file_content(self, file_path: str) -> Optional[str]:
        """
        Liest den Textinhalt einer Datei
        
        Args:
            file_path: Pfad zur Datei
            
        Returns:
            Textinhalt der Datei oder None bei Fehlern
        """
        try:
            path = Path(file_path)
            
            if not path.exists():
                logger.error(f"Datei nicht gefunden: {file_path}")
                return None
            
            # Je nach Dateityp unterschiedlich verarbeiten
            if path.suffix.lower() == '.txt':
                with open(path, 'r', encoding='utf-8') as file:
                    return file.read()
            elif path.suffix.lower() == '.pdf':
                return self._extract_pdf_text(path)
            elif path.suffix.lower() in ['.doc', '.docx']:
                return self._extract_word_text(path)
            else:
                logger.warning(f"Textextraktion für {path.suffix} nicht unterstützt")
                return None
                
        except Exception as e:
            logger.error(f"Fehler beim Lesen der Datei {file_path}: {str(e)}")
            return None
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extrahiert Text aus PDF-Dateien"""
        try:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF-Textextraktion fehlgeschlagen: {str(e)}")
            return ""
    
    def _extract_word_text(self, file_path: Path) -> str:
        """Extrahiert Text aus Word-Dokumenten"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Word-Textextraktion fehlgeschlagen: {str(e)}")
            return ""
    
    def delete_file(self, file_path: str) -> bool:
        """
        Löscht eine Datei
        
        Args:
            file_path: Pfad zur zu löschenden Datei
            
        Returns:
            True bei Erfolg, False bei Fehlern
        """
        try:
            path = Path(file_path)
            if path.exists():
                path.unlink()
                logger.info(f"Datei gelöscht: {file_path}")
                return True
            else:
                logger.warning(f"Datei zum Löschen nicht gefunden: {file_path}")
                return False
                
        except Exception as e:
            logger.error(f"Fehler beim Löschen der Datei {file_path}: {str(e)}")
            return False
    
    def cleanup_old_files(self, max_age_days: int = 7) -> int:
        """
        Löscht alte Dateien
        
        Args:
            max_age_days: Maximales Alter der Dateien in Tagen
            
        Returns:
            Anzahl der gelöschten Dateien
        """
        import time
        
        deleted_count = 0
        max_age_seconds = max_age_days * 24 * 60 * 60
        current_time = time.time()
        
        try:
            for file_path in self.upload_folder.rglob('*'):
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > max_age_seconds:
                        file_path.unlink()
                        deleted_count += 1
                        logger.info(f"Alte Datei gelöscht: {file_path}")
            
            logger.info(f"Cleanup abgeschlossen: {deleted_count} Dateien gelöscht")
            return deleted_count
            
        except Exception as e:
            logger.error(f"Fehler beim Cleanup: {str(e)}")
            return deleted_count 